import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="none"/>\n'
        f'  <w:left w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx(output_path):
    doc = docx.Document()
    
    # Page Setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Styles
    PRIMARY = RGBColor(30, 58, 138)     # Deep Blue #1E3A8A
    SECONDARY = RGBColor(37, 99, 235)   # Bright Blue #2563EB
    DARK_TEXT = RGBColor(31, 41, 55)    # Dark Charcoal #1F2937

    # Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = DARK_TEXT

    # Header / Title Block
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(4)
    run_title = title_p.add_run("แผนการดำเนินงาน (Implementation Plan)")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY

    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(18)
    run_sub = subtitle_p.add_run("ระบบลงทะเบียนผ่าน GAS และการเชื่อมต่อ Event Platform (Per-Gate Running Number)")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = SECONDARY

    # Horizontal Divider Line
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_after = Pt(18)
    p_line_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="1E3A8A"/></w:pBdr>')
    p_line._p.get_or_add_pPr().append(p_line_border)

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SECONDARY
        return h

    def add_paragraph(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
            r_bold.font.color.rgb = DARK_TEXT
        r_text = p.add_run(text)
        r_text.font.color.rgb = DARK_TEXT
        return p

    def add_bullet(text, bold_prefix="", level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.bold = True
            r_bold.font.color.rgb = DARK_TEXT
        r_text = p.add_run(text)
        r_text.font.color.rgb = DARK_TEXT
        return p

    def add_callout(title, text, bg_color="EFF6FF", border_color="3B82F6"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>\n'
            f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>\n'
            f'  <w:top w:val="none"/>\n'
            f'  <w:bottom w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)

        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r_title = p.add_run(f"📌 {title}\n")
        r_title.bold = True
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(30, 58, 138)
        
        r_text = p.add_run(text)
        r_text.font.size = Pt(10.5)
        r_text.font.color.rgb = DARK_TEXT
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 1: Overview
    add_heading_1("1. ภาพรวมและข้อกำหนดระบบ (System Overview & Key Updates)")
    add_paragraph("เอกสารนี้ระบุรายละเอียดการพัฒนาระบบลงทะเบียนผ่าน Google Apps Script (GAS) ร่วมกับ Google Sheet และเชื่อมต่อไปยัง Event Platform สำหรับสแกน Check-in หน้างานและการจับรางวัล Lucky Draw")

    add_callout(
        "ข้อกำหนดสำคัญที่ได้รับปรับปรุง (Updated Requirements)",
        "1. การออก Running Number แยกตามจุด Check-in / Gate: สามารถกำหนด Config รูปแบบ Running Number สำหรับแต่ละทางเข้าได้โดยไม่ซ้ำกัน เช่น Gate 1 ออกเลข A00001-Annnnn, Gate 2 ออกเลข B00001-Bnnnnn\n"
        "2. การนำเข้าข้อมูล (Method 1: Manual Import): Admin ดาวน์โหลดไฟล์ CSV/Excel จากแท็บ 'apploval' ของ Google Sheet แล้วนำมาอัปโหลด/Import ผ่านหน้า Web Admin ของ Event Platform"
    )

    # Section 2: Workflow 7 Steps
    add_heading_1("2. ขั้นตอนการทำงาน 7 ขั้นตอน (7-Step Workflow)")
    
    table_steps = doc.add_table(rows=8, cols=3)
    table_steps.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table_steps)
    
    headers = ["ลำดับ", "ขั้นตอน (Workflow Step)", "รายละเอียดการทำงาน"]
    hdr_cells = table_steps.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    steps_data = [
        ("1", "Admin เปิดลงทะเบียน", "Admin เปิดระบบรับลงทะเบียนบน GAS Web App"),
        ("2", "User ลงทะเบียน & รับ QR Code", "ผู้ใช้ลงทะเบียน เลือกหน่วยงานจากแท็บ 'club' ระบบสร้าง QR Code พร้อมเตือน 'โปรดบันทึก QR Code สำหรับติดตามสถานะ และ check in เข้าร่วมงาน' และบันทึกข้อมูลลงแท็บ 'registered' (สถานะ: รอการอนุมัติ)"),
        ("3", "Admin ปิดรับลงทะเบียน", "Admin ปิดรับลงทะเบียนบน GAS Web App เมนูลงทะเบียนเปลี่ยนเป็นแสดงข้อความ 'ปิดรับลงทะเบียน'"),
        ("4", "Admin อัปเดตผลอนุมัติ", "Admin ตรวจสอบรายชื่อใน Google Sheet และปรับสถานะเป็น 'อนุมัติ' (ย้าย/จัดเก็บลงแท็บ 'apploval')"),
        ("5", "Manual Import เข้า Event Platform", "Admin ดาวน์โหลดไฟล์ CSV/Excel จากแท็บ 'apploval' แล้วนำมากด Import ผ่านหน้า Web Admin ของ Event Platform"),
        ("6", "User Check-in & ออก Running Number ตาม Gate", "ผู้ใช้สแกน QR Code เพื่อ Check-in ที่หน้างานตาม Gate ที่เข้า เช่น:\n- Gate 1 Config (Prefix: A) -> ออก Running Number A00001, A00002...\n- Gate 2 Config (Prefix: B) -> ออก Running Number B00001, B00002..."),
        ("7", "Lucky Draw & ประกาศผลรางวัล", "Admin ทำการสุ่มจับรางวัล Lucky Draw บน Event Platform ด้วย Running Number ระบบบันทึกรายชื่อผู้โชคดีลงแท็บ 'winer' บน Google Sheet และแสดงในเมนู 'ประกาศผลรางวัล'")
    ]

    for row_idx, data in enumerate(steps_data, start=1):
        row_cells = table_steps.rows[row_idx].cells
        bg_color = "F9FAFB" if row_idx % 2 == 0 else "FFFFFF"
        for col_idx, cell_value in enumerate(data):
            row_cells[col_idx].text = cell_value
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            if col_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Section 3: Technical Configuration
    add_heading_1("3. การปรับปรุงทางเทคนิค (Technical Implementation)")

    add_heading_2("3.1 การตั้งค่า Running Number แยกตามจุด Check-in (Checkin Point Config)")
    add_paragraph("ในฐานข้อมูล Prisma (`prisma/schema.prisma`) จะทำการเพิ่มฟิลด์ใน `CheckinPoint` เพื่อให้ Admin สามารถกำหนด Prefix และจำนวนหลักสำหรับแต่ละทางเข้าได้อย่างอิสระ:")
    add_bullet("CheckinPoint Model: เพิ่ม prefix (เช่น 'A', 'B'), numberPadding (เช่น 5 หลัก) และ currentSeq", "Schema Update - ")
    add_bullet("Registration Model: เพิ่ม runningNumber (เช่น 'A00001', 'B00001') สำหรับบันทึกหมายเลขที่ออกให้ผู้เข้าร่วมงาน", "Registration Schema - ")

    add_callout(
        "การทำงานของ Check-in API",
        "เมื่อมีการสแกน QR Code ที่ Gate ใดก็ตาม:\n"
        "1. API จะอ่านค่า Prefix ของ Gate นั้นๆ (เช่น Gate 1 = 'A')\n"
        "2. คำนวณหาลำดับถัดไป: nextSeq = currentSeq + 1\n"
        "3. สร้าง Running Number: runningNumber = 'A' + padZero(nextSeq, 5) -> 'A00001'\n"
        "4. บันทึกค่า และอัปเดตสถานะการลงทะเบียนเป็น 'CHECKED_IN'",
        bg_color="F8FAFC", border_color="64748B"
    )

    add_heading_2("3.2 โครงสร้าง Google Sheet 4 แท็บ และ GAS Web App")
    add_bullet(" registered (รอการอนุมัติ), apploval (อนุมัติแล้ว), club (Master Data สำหรับ Dropdown หน่วยงาน), winer (รายชื่อผู้ได้รับรางวัล)", "Google Sheet (4 แท็บ): ")
    add_bullet("1. ลงทะเบียน (มีสวิตช์ปิด/เปิด, แจ้งเตือนบันทึก QR Code), 2. ติดตามสถานะ (ค้นหาตามเบอร์/อีเมล/QR Code), 3. ประกาศผลรางวัล (แสดงผลผู้โชคดี)", "GAS Web App (3 เมนู): ")

    # Section 4: Verification & Testing
    add_heading_1("4. แผนการทดสอบและการตรวจรับ (Verification Plan)")
    add_bullet("ตั้งค่า Gate 1 (Prefix A) และ Gate 2 (Prefix B) ใน Web Admin", "1. ทดสอบ Config Gate: ")
    add_bullet("Export ไฟล์ XLSX/CSV จากแท็บ apploval และ Import เข้า Event Platform", "2. ทดสอบ Manual Import: ")
    add_bullet("สแกนที่ Gate 1 ต้องได้ A00001, A00002... และสแกนที่ Gate 2 ต้องได้ B00001, B00002... (ห้ามซ้ำกันข้าม Gate)", "3. ทดสอบ Check-in & Running Number: ")
    add_bullet("สุ่มรางวัลบน web-luckydraw และย้อนกลับมาตรวจสอบความถูกต้องที่แท็บ winer บน Google Sheet", "4. ทดสอบ Lucky Draw & Winner Sync: ")

    doc.save(output_path)
    print(f"Successfully updated docx at {output_path}")

if __name__ == "__main__":
    build_docx(r"d:\Project\event-platform\Implementation_Plan_GAS_Event_Platform.docx")
